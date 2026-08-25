"""
Genera el Formato de Salida por Venta de Activo (baja de inventario + venta)
de SelectShop MB SA de CV — mismo criterio visual que generar_manual.py y
el "Formato de Salida de Equipos" (envíos) del sistema: naranja #E8431A +
negro, cajas con borde naranja, firmas al pie.

Es una PLANTILLA en blanco para imprimir y llenar a mano (o adaptar luego a
un flujo dentro del sistema), no un documento generado por instancia.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores (branding SelectShop, ver CLAUDE.md/README.md del repo) ──────
NARANJA   = RGBColor(0xE8, 0x43, 0x1A)
NEGRO     = RGBColor(0x1A, 0x1A, 0x1A)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_OSC  = RGBColor(0x44, 0x44, 0x44)
GRIS_CLAR = RGBColor(0xF5, 0xF5, 0xF5)
GRIS_MED  = RGBColor(0x88, 0x88, 0x88)

TOTAL_W = 18.0  # ancho útil de las tablas (cm) — cabe con margen de sobra
                # dentro del área imprimible (ver márgenes de página abajo)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, sz='6', color='E8431A', val='single'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), val)
        tag.set(qn('w:sz'), sz)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def no_space(para, before=0, after=4):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    pPr.append(sp)


def set_col_widths(table, widths_cm):
    """Fija el ancho por COLUMNA (no por celda suelta) — `table.columns[i].width`
    es la única forma confiable en python-docx de que Word respete el ancho:
    actualiza el `tblGrid` y las celdas de esa columna en todas las filas a
    la vez. Fijar `cell.width` celda por celda (lo que hacía la versión
    anterior de este script) deja el `tblGrid` desactualizado — Word
    reconcilia como puede y el resultado es una tabla más ancha que la
    página, recortando el borde izquierdo de la primera columna (bug real
    reportado: las etiquetas largas se veían cortadas por la izquierda,
    p. ej. "Tipo de activo:" mostrando solo "activo:")."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)


def section_title(doc, text):
    p = doc.add_paragraph()
    no_space(p, before=10, after=4)
    run = p.add_run(f'  {text}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = BLANCO
    # banda naranja de fondo, igual que sectionBand() en pdfBranding.js
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8431A')
    pPr.append(shd)
    return p


def field_row(doc, pairs, label_w=4.6):
    """Una fila de 1 o 2 pares Label/línea-en-blanco, con borde gris fino.
    `label_w` es el ancho de CADA columna de etiqueta; el resto del ancho
    disponible (TOTAL_W entre el número de pares) se reparte a la columna
    de valor de cada par."""
    table = doc.add_table(rows=1, cols=len(pairs) * 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    value_w = (TOTAL_W / len(pairs)) - label_w
    col_widths = []
    for _ in pairs:
        col_widths += [label_w, value_w]
    set_col_widths(table, col_widths)
    for i, (label, value) in enumerate(pairs):
        lc = table.rows[0].cells[i * 2]
        vc = table.rows[0].cells[i * 2 + 1]
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = GRIS_OSC
        no_space(lp)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value or '')
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = NEGRO
        no_space(vp)
        set_cell_bg(lc, 'F5F5F5')
        set_cell_border(lc, sz='4', color='CCCCCC')
        set_cell_border(vc, sz='4', color='CCCCCC')
    return table


def checkbox_line(doc, items, note=None):
    p = doc.add_paragraph()
    no_space(p, before=4, after=4)
    r = p.add_run('    '.join(f'[  ] {t}' for t in items))
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = NEGRO
    if note:
        p2 = doc.add_paragraph()
        no_space(p2, before=0, after=6)
        r2 = p2.add_run(note)
        r2.font.size = Pt(8)
        r2.font.italic = True
        r2.font.color.rgb = GRIS_MED
    return p


def signature_row(doc, labels):
    table = doc.add_table(rows=2, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(table, [TOTAL_W / len(labels)] * len(labels))
    for i, label in enumerate(labels):
        line_cell = table.rows[0].cells[i]
        line_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p = line_cell.paragraphs[0]
        no_space(p, before=30, after=2)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A1A1A')
        pbdr.append(bottom)
        pPr.append(pbdr)

        label_cell = table.rows[1].cells[i]
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        no_space(lp, before=2, after=0)
        lr = lp.add_run(label)
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = NEGRO
        sub = label_cell.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        no_space(sub, before=0, after=0)
        sr = sub.add_run('Nombre y firma')
        sr.font.size = Pt(6.5)
        sr.font.color.rgb = GRIS_MED
    return table


# ══════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
# Área imprimible con margen de 1.8cm por lado en carta (21.59cm de ancho):
# 21.59 - 1.8 - 1.8 = 17.99cm ≈ TOTAL_W (18.0cm) de las tablas — casi al
# límite. Se deja explícito aquí para que si alguien cambia TOTAL_W o los
# márgenes, no se les olvide que deben cuadrar entre sí.

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Normal'].font.color.rgb = GRIS_OSC

# ── Encabezado ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p, before=0, after=2)
run = p.add_run('FORMATO DE SALIDA POR VENTA DE ACTIVO')
run.font.size = Pt(15)
run.font.bold = True
run.font.color.rgb = NARANJA

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p, before=0, after=6)
run = p.add_run('Baja de inventario por venta — SelectShop MB SA de CV — Sistemas IT & BI')
run.font.size = Pt(9)
run.font.color.rgb = GRIS_MED

# línea naranja
p = doc.add_paragraph()
no_space(p, before=0, after=8)
pPr = p._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '18')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'E8431A')
pbdr.append(bottom)
pPr.append(pbdr)

field_row(doc, [('FOLIO:', ''), ('FECHA:', '')])

# ── Datos del activo ─────────────────────────────────────────────────
section_title(doc, 'DATOS DEL ACTIVO')
field_row(doc, [('Tipo de activo:', ''), ('Marca / Modelo:', '')])
field_row(doc, [('No. de serie:', ''), ('No. de inventario:', '')])
field_row(doc, [('Sucursal / resguardo actual:', '')])
checkbox_line(
    doc,
    ['Bueno', 'Regular', 'Dañado / con detalle (especificar abajo)'],
    note='Condición física y funcional del equipo al momento de la venta.',
)
p = doc.add_paragraph()
no_space(p, before=0, after=10)
pPr = p._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '4')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'CCCCCC')
pbdr.append(bottom)
pPr.append(pbdr)
r = p.add_run('Detalle de la condición / observaciones: _______________________________________________')
r.font.size = Pt(8.5)
r.font.color.rgb = GRIS_OSC

checkbox_line(
    doc,
    ['Equipo formateado / restablecido a valores de fábrica'],
    note='Se retiraron cuentas y datos corporativos de SelectShop (correo, VPN, licencias) antes de la entrega.',
)

# ── Motivo de baja ───────────────────────────────────────────────────
section_title(doc, 'MOTIVO DE BAJA')
p = doc.add_paragraph()
no_space(p, before=4, after=6)
r = p.add_run('[ X ]  VENTA DEL ACTIVO')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = NEGRO
p2 = doc.add_paragraph()
no_space(p2, before=0, after=8)
r2 = p2.add_run(
    'Se da de baja del inventario de activos de SelectShop MB SA de CV el equipo '
    'descrito arriba, por concepto de venta, a partir de la fecha indicada en este documento.'
)
r2.font.size = Pt(9)
r2.font.color.rgb = GRIS_OSC

# ── Datos del comprador ──────────────────────────────────────────────
section_title(doc, 'DATOS DEL COMPRADOR')
checkbox_line(doc, ['Empleado de SelectShop', 'Externo / tercero'])
field_row(doc, [('Nombre completo:', '')])
field_row(doc, [('No. de empleado (si aplica):', ''), ('Puesto / Sucursal (si aplica):', '')])
field_row(doc, [('Identificación oficial No. (si es externo):', ''), ('Teléfono:', '')])
field_row(doc, [('Domicilio (si es externo):', '')])

# ── Datos de la venta ─────────────────────────────────────────────────
section_title(doc, 'DATOS DE LA VENTA')
field_row(doc, [('Monto de venta: $', ''), ('Fecha de pago:', '')])
checkbox_line(doc, ['Efectivo', 'Transferencia', 'Descuento vía nómina', 'Otro:'])
field_row(doc, [('Referencia / comprobante (opcional):', '')])

# ── Condiciones de la venta ─────────────────────────────────────────
section_title(doc, 'CONDICIONES DE LA VENTA')
p = doc.add_paragraph()
no_space(p, before=4, after=8)
r = p.add_run(
    'El comprador declara haber revisado el activo descrito en este documento y acepta '
    'recibirlo en el estado físico y funcional en que se encuentra a la fecha de esta '
    'venta ("tal cual"), sin garantía de ningún tipo por parte de SelectShop MB SA de CV. '
    'A partir de la firma de este documento, el activo deja de formar parte del inventario '
    'y del resguardo de la empresa, y toda responsabilidad por su uso, mantenimiento o '
    'reparación posterior corresponde exclusivamente al comprador.'
)
r.font.size = Pt(8.5)
r.font.color.rgb = GRIS_OSC
r.font.italic = True

# ── Firmas ────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p, before=16, after=14)
r = p.add_run('FIRMAS Y AUTORIZACIONES')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = NEGRO

signature_row(doc, ['Entrega — IT', 'Autoriza — Gerente de Sistemas', 'Compra — recibí el activo de conformidad'])

# ── Pie ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p, before=14, after=0)
r = p.add_run('SELECTSHOP MB  |  Sistemas IT & BI  |  Uso Interno')
r.font.size = Pt(7)
r.font.color.rgb = GRIS_MED

output = '/Users/systems/Projects/assets-manager/responsiva_ref/Formato_Salida_Venta_Activo.docx'
import os
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print('Guardado en', output)
