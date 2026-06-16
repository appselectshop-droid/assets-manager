"""
Genera el manual de usuario del Sistema de Gestión de Activos
Colores SelectShop: naranja #E8431A, negro #111111
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ────────────────────────────────────────────
NARANJA   = RGBColor(0xE8, 0x43, 0x1A)
NEGRO     = RGBColor(0x11, 0x11, 0x11)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_OSC  = RGBColor(0x44, 0x44, 0x44)
GRIS_CLAR = RGBColor(0xF5, 0xF5, 0xF5)
GRIS_MED  = RGBColor(0x88, 0x88, 0x88)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'E8431A'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '0')
    pPr.append(sp)

def add_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    no_space_para(p)

def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml  import OxmlElement as _el
    br = _el('w:br')
    br.set(_qn('w:type'), 'page')
    return br

# ────────────────────────────────────────────────────────
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
doc.styles['Normal'].font.name  = 'Calibri'
doc.styles['Normal'].font.size  = Pt(11)
doc.styles['Normal'].font.color.rgb = GRIS_OSC

# ══════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════
# Espacio superior
for _ in range(6):
    p = doc.add_paragraph()
    no_space_para(p)

# Línea naranja decorativa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.color.rgb = NARANJA
run.font.size      = Pt(14)
no_space_para(p)

p = doc.add_paragraph()
no_space_para(p)

# Título principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Gestión de Activos')
run.font.name      = 'Calibri'
run.font.size      = Pt(32)
run.font.bold      = True
run.font.color.rgb = NEGRO
no_space_para(p)

p = doc.add_paragraph()
no_space_para(p)

# Subtítulo naranja
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE USUARIO')
run.font.name      = 'Calibri'
run.font.size      = Pt(18)
run.font.bold      = True
run.font.color.rgb = NARANJA
no_space_para(p)

p = doc.add_paragraph()
no_space_para(p)

# Línea naranja decorativa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.color.rgb = NARANJA
run.font.size      = Pt(14)
no_space_para(p)

for _ in range(4):
    p = doc.add_paragraph()
    no_space_para(p)

# Versión y fecha
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.0  ·  Junio 2026')
run.font.size      = Pt(12)
run.font.color.rgb = GRIS_MED
no_space_para(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Select Shop')
run.font.size      = Pt(12)
run.font.bold      = True
run.font.color.rgb = GRIS_MED
no_space_para(p)

# ── Helpers para el cuerpo ─────────────────────────────
def page_break(doc):
    doc.add_page_break()

def heading1(doc, texto):
    """Encabezado de sección con línea naranja"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(20)
    run.font.bold      = True
    run.font.color.rgb = NARANJA
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '300')
    sp.set(qn('w:after'),  '80')
    pPr.append(sp)

    # Línea naranja debajo
    p2 = doc.add_paragraph()
    run2 = p2.add_run('─' * 60)
    run2.font.color.rgb = NARANJA
    run2.font.size      = Pt(9)
    pPr2 = p2._p.get_or_add_pPr()
    sp2  = OxmlElement('w:spacing')
    sp2.set(qn('w:before'), '0')
    sp2.set(qn('w:after'),  '120')
    pPr2.append(sp2)

def heading2(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = NEGRO
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '200')
    sp.set(qn('w:after'),  '80')
    pPr.append(sp)

def heading3(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(12)
    run.font.bold      = True
    run.font.color.rgb = NARANJA
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '140')
    sp.set(qn('w:after'),  '40')
    pPr.append(sp)

def body(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size      = Pt(11)
    run.font.color.rgb = GRIS_OSC
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '80')
    pPr.append(sp)
    return p

def bullet(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texto)
    run.font.size      = Pt(11)
    run.font.color.rgb = GRIS_OSC
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '60')
    pPr.append(sp)

def step(doc, numero, texto):
    """Paso numerado con número en naranja"""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{numero}.  ')
    r1.font.bold      = True
    r1.font.color.rgb = NARANJA
    r1.font.size      = Pt(11)
    r2 = p.add_run(texto)
    r2.font.color.rgb = GRIS_OSC
    r2.font.size      = Pt(11)
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '80')
    pPr.append(sp)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)

def nota(doc, texto):
    """Caja de nota con fondo naranja claro"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF3EE')
    cell.width = Cm(14)
    p = cell.paragraphs[0]
    r1 = p.add_run('📌  Nota:  ')
    r1.font.bold      = True
    r1.font.color.rgb = NARANJA
    r1.font.size      = Pt(10.5)
    r2 = p.add_run(texto)
    r2.font.color.rgb = GRIS_OSC
    r2.font.size      = Pt(10.5)
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '80')
    sp.set(qn('w:after'),  '80')
    pPr.append(sp)
    doc.add_paragraph()

def table_header(doc, columnas, anchos=None):
    """Tabla con encabezado naranja"""
    tbl = doc.add_table(rows=1, cols=len(columnas))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = tbl.rows[0].cells
    for i, col in enumerate(columnas):
        set_cell_bg(hdr_cells[i], 'E8431A')
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col)
        run.font.bold      = True
        run.font.color.rgb = BLANCO
        run.font.size      = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        no_space_para(p)
        if anchos:
            hdr_cells[i].width = Cm(anchos[i])
    return tbl

def table_row(tbl, valores, alternado=False):
    row = tbl.add_row()
    for i, val in enumerate(valores):
        if alternado:
            set_cell_bg(row.cells[i], 'F9F9F9')
        p = row.cells[i].paragraphs[0]
        run = p.add_run(str(val))
        run.font.size      = Pt(10)
        run.font.color.rgb = GRIS_OSC
        no_space_para(p)
    return row

def espacio(doc):
    p = doc.add_paragraph()
    no_space_para(p)

# ══════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '📋  Tabla de Contenido')

indice = [
    ('1.', 'Introducción',                                     '3'),
    ('2.', 'Acceso al Sistema',                                '3'),
    ('3.', 'Dashboard — Pantalla Principal',                   '4'),
    ('4.', 'Módulo de Activos',                                '6'),
    ('  4.1', 'Ver y filtrar activos',                         '6'),
    ('  4.2', 'Registrar un nuevo activo',                     '7'),
    ('  4.3', 'Editar un activo',                              '9'),
    ('  4.4', 'Asignar / devolver un activo',                  '9'),
    ('  4.5', 'Acciones en lote',                              '10'),
    ('  4.6', 'Importar activos desde Excel',                  '10'),
    ('  4.7', 'Exportar activos a Excel',                      '11'),
    ('5.', 'Módulo de Empleados',                              '11'),
    ('  5.1', 'Ver y buscar empleados',                        '11'),
    ('  5.2', 'Registrar un empleado',                         '12'),
    ('  5.3', 'Perfil de empleado y activos asignados',        '12'),
    ('6.', 'Módulo de Asignaciones',                           '13'),
    ('7.', 'Gestión de Usuarios',                              '14'),
    ('8.', 'Preguntas Frecuentes',                             '14'),
    ('9.', 'Glosario de Términos',                             '15'),
]

tbl = doc.add_table(rows=0, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for num, titulo, pag in indice:
    row = tbl.add_row()
    cells = row.cells
    r0 = cells[0].paragraphs[0].add_run(num)
    r0.font.bold      = True
    r0.font.color.rgb = NARANJA if not num.startswith('  ') else GRIS_OSC
    r0.font.size      = Pt(11)
    cells[0].width    = Cm(1.5)

    r1 = cells[1].paragraphs[0].add_run(titulo)
    r1.font.color.rgb = NEGRO
    r1.font.size      = Pt(11)
    cells[1].width    = Cm(12)

    r2 = cells[2].paragraphs[0].add_run(pag)
    r2.font.color.rgb = GRIS_MED
    r2.font.size      = Pt(11)
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cells[2].width = Cm(1.5)

    for c in cells:
        no_space_para(c.paragraphs[0])
        pPr = c.paragraphs[0]._p.get_or_add_pPr()
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '60')
        sp.set(qn('w:after'),  '60')
        pPr.append(sp)

# ══════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '1.  Introducción')
body(doc, 'El Sistema de Gestión de Activos de Select Shop es una aplicación web diseñada para llevar el control completo del inventario de equipos tecnológicos de la empresa. Permite registrar, asignar, editar y dar seguimiento a todos los activos (laptops, escritorios, celulares, periféricos, accesorios, etc.) y a los empleados a quienes están asignados.')
espacio(doc)
body(doc, 'Con este sistema podrás:')
bullet(doc, 'Consultar en tiempo real el inventario completo de activos.')
bullet(doc, 'Registrar nuevos equipos con todas sus especificaciones técnicas.')
bullet(doc, 'Asignar y regresar activos a empleados de forma sencilla.')
bullet(doc, 'Filtrar y buscar activos por tipo, estado, serie, IMEI, contrato, etc.')
bullet(doc, 'Importar y exportar inventario en archivos Excel (.xlsx).')
bullet(doc, 'Ver estadísticas y resúmenes en el Dashboard interactivo.')

espacio(doc)
heading2(doc, 'Requisitos para acceder')
bullet(doc, 'Computadora, laptop o tablet con acceso a internet.')
bullet(doc, 'Navegador web actualizado (Chrome, Edge, Firefox, Safari).')
bullet(doc, 'Credenciales de acceso (correo y contraseña) proporcionadas por el administrador.')

# ══════════════════════════════════════════════════════════
# 2. ACCESO AL SISTEMA
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '2.  Acceso al Sistema')

heading2(doc, '2.1  Iniciar sesión')
body(doc, 'Para acceder al sistema sigue estos pasos:')
step(doc, 1, 'Abre tu navegador web e ingresa la dirección URL del sistema.')
step(doc, 2, 'En la pantalla de inicio de sesión, escribe tu correo electrónico institucional.')
step(doc, 3, 'Ingresa tu contraseña.')
step(doc, 4, 'Haz clic en el botón "Iniciar sesión".')
step(doc, 5, 'Serás redirigido automáticamente al Dashboard principal.')

espacio(doc)
nota(doc, 'Si olvidaste tu contraseña o no tienes acceso, comunícate con el administrador del sistema para restablecer tus credenciales.')

heading2(doc, '2.2  Cerrar sesión')
body(doc, 'Para cerrar sesión de forma segura, haz clic en el ícono de usuario ubicado en la parte superior del menú lateral y selecciona "Cerrar sesión". Siempre cierra sesión cuando termines de usar el sistema, especialmente en equipos compartidos.')

# ══════════════════════════════════════════════════════════
# 3. DASHBOARD
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '3.  Dashboard — Pantalla Principal')
body(doc, 'El Dashboard es la primera pantalla que verás al iniciar sesión. Muestra un resumen visual del estado del inventario y las actividades recientes. Se actualiza automáticamente cada vez que entras a la sección.')

heading2(doc, '3.1  Indicadores principales (KPIs)')
body(doc, 'En la parte superior encontrarás 5 tarjetas con los indicadores clave:')

tbl = table_header(doc, ['Indicador', 'Descripción'], [5, 10])
filas_kpi = [
    ('👥 Empleados',       'Total de empleados registrados en el sistema.'),
    ('💻 Activos totales', 'Número total de activos en el inventario.'),
    ('🔗 Asignados',       'Activos actualmente asignados a algún empleado.'),
    ('✅ Disponibles',      'Activos libres, sin asignar.'),
    ('🚫 De baja',         'Activos retirados del servicio.'),
]
for i, (ind, desc) in enumerate(filas_kpi):
    table_row(tbl, [ind, desc], alternado=(i % 2 == 1))
espacio(doc)

body(doc, 'Puedes hacer clic en cualquier tarjeta para ir directamente al módulo correspondiente.')

heading2(doc, '3.2  Filtros por sucursal y departamento')
body(doc, 'Debajo de los KPIs encontrarás chips (botones de filtro) para ver los datos por sucursal y/o departamento. Al activar un filtro, todos los indicadores y gráficas del Dashboard se actualizan para mostrar solo los datos del grupo seleccionado.')
bullet(doc, 'Haz clic en una sucursal para filtrar por ella.')
bullet(doc, 'Opcionalmente, selecciona también un departamento dentro de esa sucursal.')
bullet(doc, 'Haz clic en "✕ Limpiar filtros" para volver a la vista global.')

heading2(doc, '3.3  Activos por categoría')
body(doc, 'La tarjeta "Activos por categoría" muestra cuántos activos hay en cada grupo:')
bullet(doc, '💻 Cómputo — Laptops, Escritorios, All-in-One')
bullet(doc, '📱 Móviles — Celulares, Tablets, Cargadores de celular')
bullet(doc, '🖱️ Periféricos — Monitores, Mouse, Teclados, Cargadores de laptop')
bullet(doc, '📦 Otros — Accesorios y otros equipos')
espacio(doc)
body(doc, 'Puedes hacer clic en cualquier categoría para ver el desglose por tipo de equipo (ejemplo: dentro de Cómputo verás cuántos son Laptop, cuántos Escritorio y cuántos All-in-One). Usa el botón "←" para regresar a la vista de categorías.')

heading2(doc, '3.4  Estado del inventario')
body(doc, 'La gráfica circular (donut) muestra la proporción del inventario total por estado: Asignados (naranja), Disponibles (verde) y De baja (rojo). Los datos siempre son globales, independientemente del filtro activo.')

heading2(doc, '3.5  Propiedad — Cómputo')
body(doc, 'Esta tarjeta muestra cuántos equipos de cómputo (laptops, escritorios y all-in-ones) son de arrendamiento y cuántos son propiedad de la empresa. Incluye el desglose por tipo de equipo.')
bullet(doc, '🟠 Arrendamiento — Equipos en renta o leasing.')
bullet(doc, '🔵 Propia — Equipos propiedad de la empresa.')
bullet(doc, 'Sin definir — Equipos a los que no se les ha asignado tipo de propiedad.')

heading2(doc, '3.6  Últimas asignaciones y Top empleados')
body(doc, '"Últimas asignaciones" muestra las 6 asignaciones más recientes con el nombre del empleado y el equipo. "Top empleados" lista a los 5 empleados con más activos asignados. Ambas secciones respetan el filtro de sucursal/departamento activo.')

# ══════════════════════════════════════════════════════════
# 4. MÓDULO DE ACTIVOS
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '4.  Módulo de Activos')
body(doc, 'El módulo de Activos es el núcleo del sistema. Aquí puedes ver, registrar, editar, asignar, importar y exportar todos los activos del inventario. Accede desde el menú lateral haciendo clic en "Activos".')

heading2(doc, '4.1  Ver y filtrar activos')
body(doc, 'Al entrar al módulo verás una tabla con todos los activos. En la parte superior hay pestañas para filtrar por categoría:')

tbl = table_header(doc, ['Pestaña', 'Tipos de activos incluidos'], [4, 11])
table_row(tbl, ['📋 Todos',              'Muestra todos los activos registrados.'],         False)
table_row(tbl, ['💻 Equipo de cómputo',  'Laptops, Escritorios, All-in-One.'],              True)
table_row(tbl, ['📱 Celulares',          'Celulares, Tablets, Cargadores de celular.'],      False)
table_row(tbl, ['🖱️ Periféricos',        'Monitores, Mouse, Teclados, Cargadores laptop.'], True)
table_row(tbl, ['📦 Accesorios / Otros', 'Accesorios y equipos varios.'],                   False)
espacio(doc)

body(doc, 'Además, puedes usar la barra de herramientas para:')
bullet(doc, 'Buscar por cualquier campo: marca, modelo, número de serie, IMEI, número de línea, número de contrato, etc.')
bullet(doc, 'Filtrar por estado: Disponible, Asignado o De baja.')
espacio(doc)
nota(doc, 'Si hay activos con número de serie o número de línea duplicados, aparecerá una alerta naranja en la parte superior de la tabla indicando los duplicados detectados.')

heading2(doc, '4.2  Registrar un nuevo activo')
body(doc, 'Para agregar un activo al inventario:')
step(doc, 1, 'Haz clic en el botón "+ Registrar activo" en la esquina superior derecha.')
step(doc, 2, 'Se abrirá un formulario. Selecciona el tipo de activo (Laptop, Escritorio, Celular, etc.).')
step(doc, 3, 'Llena los datos generales: Marca, Modelo, Número de serie, Etiqueta de inventario, Estado y Fecha de compra.')
step(doc, 4, 'Completa las especificaciones técnicas específicas del tipo de equipo (ver tabla abajo).')
step(doc, 5, 'Si el activo ya tiene dueño, activa la opción "Asignar a un empleado ahora" y busca al empleado.')
step(doc, 6, 'Haz clic en "Registrar activo" (o "Registrar y asignar" si elegiste un empleado).')

espacio(doc)
body(doc, 'Las especificaciones varían según el tipo de activo:')

tbl = table_header(doc, ['Tipo', 'Especificaciones disponibles'], [4, 11])
specs = [
    ('Laptop / Escritorio\n/ All-in-One', 'Propiedad (Propia/Arrendamiento), Nº contrato, AnyDesk ID, Procesador, RAM, Almacenamiento, S.O., Color, Nº serie cargador, Accesorios incluidos (cargador, monitor, mouse, teclado).'),
    ('Celular / Tablet',     'Nº contrato, Razón social, Correo Gmail, IMEI 1, IMEI 2, Nº de línea, Operadora, Costo del plan, Almacenamiento, RAM, S.O., Color.'),
    ('Monitor',              'Tamaño, Resolución, Tipo de panel, Frecuencia, Tipo de conexión.'),
    ('Mouse / Teclado',      'Tipo de conexión, Distribución (teclado), Color.'),
    ('Cargador',             'Potencia (Watts), Tipo de conector, Marca/modelo compatible.'),
    ('Accesorio / Otro',     'Tipo, Tipo de conexión, Color, Descripción adicional.'),
]
for i, (tipo, spec) in enumerate(specs):
    table_row(tbl, [tipo, spec], alternado=(i % 2 == 1))
espacio(doc)

nota(doc, 'El sistema detecta automáticamente si el número de serie o número de línea ya existe en otro activo y muestra una advertencia antes de guardar.')

heading2(doc, '4.3  Editar un activo')
body(doc, 'Para modificar la información de un activo existente:')
step(doc, 1, 'Ubica el activo en la tabla (usa la búsqueda si es necesario).')
step(doc, 2, 'Haz clic en el botón "Editar" al final de su fila.')
step(doc, 3, 'El formulario se abrirá con los datos actuales del activo.')
step(doc, 4, 'Modifica los campos que necesites.')
step(doc, 5, 'Haz clic en "Guardar cambios".')

espacio(doc)
nota(doc, 'Si el activo está asignado, en la sección inferior del formulario aparecerá la asignación actual. Puedes cambiar de empleado o devolver el activo desde ahí.')

heading2(doc, '4.4  Asignar y devolver activos')
heading3(doc, 'Asignar desde el módulo de Activos')
step(doc, 1, 'Abre el formulario de edición del activo.')
step(doc, 2, 'En la sección "Asignación actual" (si ya está asignado) haz clic en "Cambiar".')
step(doc, 3, 'Si no está asignado, activa la casilla "Asignar a un empleado ahora".')
step(doc, 4, 'Escribe el nombre o número del empleado en el buscador.')
step(doc, 5, 'Selecciona al empleado de la lista y agrega notas de asignación si deseas.')
step(doc, 6, 'Haz clic en "Guardar y asignar".')

espacio(doc)
heading3(doc, 'Devolver un activo')
step(doc, 1, 'Abre el formulario de edición del activo asignado.')
step(doc, 2, 'En la sección de asignación, haz clic en "Devolver".')
step(doc, 3, 'Haz clic en "Guardar y devolver". El activo quedará como "Disponible".')

heading2(doc, '4.5  Acciones en lote')
body(doc, 'Puedes seleccionar varios activos a la vez para aplicar acciones de forma masiva:')
step(doc, 1, 'Marca las casillas de los activos que deseas seleccionar (o usa la casilla del encabezado para seleccionar todos los visibles).')
step(doc, 2, 'Aparecerá una barra inferior con las acciones disponibles.')
step(doc, 3, 'Elige la acción deseada:')
bullet(doc, '✅ Marcar disponible — Cambia el estado a "Disponible".')
bullet(doc, '🚫 Dar de baja — Cambia el estado a "De baja".')
bullet(doc, '🗑️ Eliminar — Elimina permanentemente los activos seleccionados.')
espacio(doc)
nota(doc, 'La eliminación es irreversible. Asegúrate de seleccionar solo los activos correctos antes de confirmar.')

heading2(doc, '4.6  Importar activos desde Excel')
body(doc, 'El sistema permite importar activos masivamente desde un archivo Excel (.xlsx):')
step(doc, 1, 'Haz clic en el botón "📥 Importar Excel ▾" y selecciona la categoría a importar (Cómputo, Celulares, Periféricos o Accesorios).')
step(doc, 2, 'Descarga la plantilla de ejemplo que aparece en el modal para ver el formato requerido.')
step(doc, 3, 'Llena la plantilla con los datos de tus activos. Respeta los nombres de las columnas exactamente.')
step(doc, 4, 'Arrastra el archivo terminado al área marcada o haz clic en "Seleccionar archivo".')
step(doc, 5, 'El sistema mostrará una vista previa de los datos. Revisa que todo esté correcto.')
step(doc, 6, 'Haz clic en "Importar" para cargar todos los registros.')

espacio(doc)
nota(doc, 'Si alguna fila tiene errores de formato o datos faltantes, el sistema te indicará cuáles filas no pudieron importarse y el motivo.')

heading2(doc, '4.7  Exportar activos a Excel')
step(doc, 1, 'Haz clic en "📤 Exportar Excel ▾".')
step(doc, 2, 'Selecciona la categoría que deseas exportar (Todos, Cómputo, Celulares, Periféricos o Accesorios).')
step(doc, 3, 'Se descargará automáticamente un archivo .xlsx con todos los activos de esa categoría y sus especificaciones.')

# ══════════════════════════════════════════════════════════
# 5. MÓDULO DE EMPLEADOS
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '5.  Módulo de Empleados')
body(doc, 'El módulo de Empleados te permite gestionar el directorio de personal y ver los activos asignados a cada uno. Accede desde el menú lateral haciendo clic en "Empleados".')

heading2(doc, '5.1  Ver y buscar empleados')
body(doc, 'La pantalla principal muestra la lista de todos los empleados registrados. Puedes:')
bullet(doc, 'Usar la barra de búsqueda para encontrar empleados por nombre, número de empleado o departamento.')
bullet(doc, 'Filtrar por sucursal o departamento usando los botones desplegables en la parte superior.')
bullet(doc, 'Ver cuántos activos tiene asignados cada empleado directamente en la tabla.')

heading2(doc, '5.2  Registrar un empleado')
step(doc, 1, 'Haz clic en "+ Registrar empleado".')
step(doc, 2, 'Llena el formulario con los datos del empleado:')
bullet(doc, 'Nombre completo (obligatorio).')
bullet(doc, 'Número de empleado (obligatorio).')
bullet(doc, 'Razón social, Sucursal, Puesto, Área, Departamento.')
bullet(doc, 'Correos corporativos y cuentas Gmail (puedes agregar varios).')
step(doc, 3, 'Haz clic en "Registrar empleado" para guardar.')

heading2(doc, '5.3  Perfil de empleado y activos asignados')
body(doc, 'Al hacer clic en el nombre de un empleado (o en el botón "Ver detalle"), accedes a su perfil completo. Aquí puedes:')
bullet(doc, 'Ver todos los activos que tiene actualmente asignados, con fecha de asignación y notas.')
bullet(doc, 'Hacer clic en "Editar" sobre un activo asignado para modificar sus datos o las notas de asignación.')
bullet(doc, 'Hacer clic en "Regresar" para devolver un activo y dejarlo disponible.')
bullet(doc, 'Hacer clic en "+ Asignar activo" para asignarle un nuevo equipo disponible.')
espacio(doc)
heading3(doc, 'Asignar un activo desde el perfil del empleado')
step(doc, 1, 'Haz clic en "+ Asignar activo".')
step(doc, 2, 'En el modal que aparece, usa las pestañas de tipo para filtrar (Laptop, Celular, Monitor, etc.).')
step(doc, 3, 'Busca el activo por nombre, serie, IMEI, contrato u otros datos.')
step(doc, 4, 'Haz clic en el activo deseado para seleccionarlo.')
step(doc, 5, 'Agrega notas de asignación opcionales.')
step(doc, 6, 'Haz clic en "Asignar".')

espacio(doc)
nota(doc, 'Si el activo que necesitas no está registrado aún, puedes registrarlo directamente desde este modal con el botón "+ Registrar activo".')

# ══════════════════════════════════════════════════════════
# 6. MÓDULO DE ASIGNACIONES
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '6.  Módulo de Asignaciones')
body(doc, 'El módulo de Asignaciones muestra el historial completo de todas las asignaciones activas: qué activos están asignados, a quién y desde cuándo. Accede desde el menú lateral.')

heading2(doc, '6.1  Ver asignaciones')
body(doc, 'La tabla muestra todas las asignaciones activas con la siguiente información:')

tbl = table_header(doc, ['Columna', 'Descripción'], [5, 10])
tabla_asig = [
    ('Empleado',           'Nombre y número del empleado con el activo asignado.'),
    ('Activo',             'Tipo, marca y modelo del equipo asignado.'),
    ('Número de serie',    'Identificador único del activo.'),
    ('Fecha de asignación','Cuándo se realizó la asignación.'),
    ('Notas',              'Observaciones registradas al momento de asignar.'),
]
for i, (col, desc) in enumerate(tabla_asig):
    table_row(tbl, [col, desc], alternado=(i % 2 == 1))
espacio(doc)

body(doc, 'Desde esta pantalla puedes regresar activos directamente haciendo clic en "Regresar" en la fila correspondiente.')

heading2(doc, '6.2  Búsqueda en asignaciones')
body(doc, 'Usa la barra de búsqueda para encontrar asignaciones por nombre del empleado, tipo de activo, marca, modelo o número de serie. También puedes filtrar por tipo de activo usando el selector desplegable.')

# ══════════════════════════════════════════════════════════
# 7. GESTIÓN DE USUARIOS
# ══════════════════════════════════════════════════════════
heading1(doc, '7.  Gestión de Usuarios')
body(doc, 'El módulo de Usuarios es exclusivo para administradores del sistema. Permite crear y gestionar las cuentas de acceso.')

heading2(doc, '7.1  Crear un nuevo usuario')
step(doc, 1, 'Accede a "Usuarios" desde el menú lateral (solo visible para administradores).')
step(doc, 2, 'Haz clic en "+ Nuevo usuario".')
step(doc, 3, 'Ingresa el nombre, correo electrónico y contraseña.')
step(doc, 4, 'Selecciona el rol: Administrador o Usuario estándar.')
step(doc, 5, 'Haz clic en "Crear usuario".')

espacio(doc)
tbl = table_header(doc, ['Rol', 'Permisos'], [4, 11])
table_row(tbl, ['Administrador', 'Acceso completo: puede crear/eliminar usuarios, gestionar todo el inventario y configuraciones.'], False)
table_row(tbl, ['Usuario',       'Puede ver y editar activos, empleados y asignaciones, pero no gestionar cuentas de usuario.'],   True)
espacio(doc)

nota(doc, 'Solo el administrador puede crear o eliminar usuarios. Si necesitas acceso o un cambio de contraseña, contacta al administrador.')

# ══════════════════════════════════════════════════════════
# 8. PREGUNTAS FRECUENTES
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '8.  Preguntas Frecuentes')

faqs = [
    ('¿Puedo asignar el mismo activo a dos empleados?',
     'No. Un activo solo puede estar asignado a un empleado a la vez. Para reasignarlo, primero debes devolverlo o cambiarlo directamente desde el formulario de edición.'),
    ('¿Qué pasa si registro un activo con un número de serie duplicado?',
     'El sistema mostrará una advertencia antes de guardar. Si ya existen duplicados en el inventario, aparecerá una alerta en la parte superior de la tabla de activos indicando los registros afectados.'),
    ('¿Cómo sé si un activo es de arrendamiento o propio?',
     'En el formulario del activo, en la sección de especificaciones de cómputo, encontrarás el campo "Propiedad" donde puedes seleccionar "Arrendamiento" o "Propia". En el Dashboard puedes ver el resumen general en la tarjeta "Propiedad — Cómputo".'),
    ('¿Puedo recuperar un activo eliminado?',
     'No. La eliminación es permanente. Te recomendamos usar el estado "De baja" en lugar de eliminar cuando quieras retirar un equipo del servicio activo pero mantener el historial.'),
    ('¿Cómo actualizo las especificaciones de un activo (RAM, procesador, etc.)?',
     'Abre el formulario de edición del activo (botón "Editar") y modifica los campos en la sección "Especificaciones". Guarda los cambios con el botón correspondiente.'),
    ('¿Puedo exportar solo los celulares o solo las laptops?',
     'Sí. Al usar "Exportar Excel", selecciona la categoría que necesitas. Por ejemplo, "Celulares / Tablets" exportará solo ese tipo de activos con todas sus columnas específicas (IMEI, línea, operadora, etc.).'),
    ('¿El sistema funciona en celular o tablet?',
     'Sí, el sistema es responsivo y se adapta a pantallas más pequeñas. Sin embargo, para mayor comodidad se recomienda usar una pantalla de laptop o escritorio.'),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'❓  {pregunta}')
    run.font.bold      = True
    run.font.color.rgb = NEGRO
    run.font.size      = Pt(11)
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '160')
    sp.set(qn('w:after'),  '40')
    pPr.append(sp)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(respuesta)
    run2.font.color.rgb = GRIS_OSC
    run2.font.size      = Pt(11)
    pPr2 = p2._p.get_or_add_pPr()
    sp2  = OxmlElement('w:spacing')
    sp2.set(qn('w:before'), '0')
    sp2.set(qn('w:after'),  '60')
    pPr2.append(sp2)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    pPr2.append(ind)

    # línea divisoria sutil
    p3 = doc.add_paragraph()
    r3 = p3.add_run('─' * 80)
    r3.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
    r3.font.size      = Pt(7)
    no_space_para(p3)

# ══════════════════════════════════════════════════════════
# 9. GLOSARIO
# ══════════════════════════════════════════════════════════
page_break(doc)
heading1(doc, '9.  Glosario de Términos')
body(doc, 'A continuación se definen los términos técnicos y conceptos clave utilizados en el Sistema de Gestión de Activos.')
espacio(doc)

glosario = [
    ('Activo',
     'Cualquier bien tecnológico registrado en el inventario del sistema: laptops, escritorios, celulares, monitores, teclados, accesorios, etc.'),
    ('All-in-One (AIO)',
     'Computadora de escritorio en la que el monitor y la unidad central están integrados en una sola carcasa.'),
    ('AnyDesk ID',
     'Identificador único del software de acceso remoto AnyDesk instalado en un equipo. Permite la asistencia técnica a distancia.'),
    ('Arrendamiento',
     'Modalidad de propiedad en la que el equipo es alquilado o rentado a un tercero mediante un contrato. No es propiedad directa de la empresa.'),
    ('Asignación',
     'Registro que vincula un activo con un empleado específico, indicando quién tiene en uso ese equipo y desde cuándo.'),
    ('Baja',
     'Estado de un activo que ha sido retirado del servicio activo. El equipo permanece en el inventario como referencia histórica pero no está disponible para asignarse.'),
    ('Cargador celular',
     'Adaptador de corriente o cable utilizado para cargar dispositivos móviles (celulares o tablets).'),
    ('Cargador laptop',
     'Adaptador de corriente específico para computadoras portátiles.'),
    ('Dashboard',
     'Pantalla principal del sistema que muestra un resumen visual e interactivo del estado general del inventario mediante gráficas, indicadores y estadísticas.'),
    ('Disponible',
     'Estado de un activo que no está asignado a ningún empleado y puede ser utilizado o asignado en cualquier momento.'),
    ('IMEI',
     'International Mobile Equipment Identity. Número de identificación único de 15 dígitos asignado a cada dispositivo móvil (celular o tablet con SIM). Se usa para identificar el equipo de forma global.'),
    ('Importar',
     'Proceso de cargar datos masivamente al sistema desde un archivo Excel (.xlsx), siguiendo el formato de plantilla definido por el sistema.'),
    ('Inventario',
     'Conjunto total de activos registrados en el sistema, independientemente de su estado (disponibles, asignados o de baja).'),
    ('KPI',
     'Key Performance Indicator (Indicador Clave de Desempeño). En el Dashboard, son los números grandes que muestran el resumen rápido del inventario: total de activos, asignados, disponibles, etc.'),
    ('No. de contrato',
     'Número o clave del contrato de arrendamiento o servicio asociado a un equipo. Aplica principalmente a celulares (plan de datos) y equipos en renta.'),
    ('No. de línea',
     'Número telefónico asociado a un celular o tablet con plan de datos o voz.'),
    ('No. de serie',
     'Número único asignado por el fabricante para identificar un equipo específico. Distinto al número de inventario interno.'),
    ('Operadora',
     'Empresa de telecomunicaciones que provee el servicio de voz y datos a un dispositivo móvil (por ejemplo: Telcel, AT&T, Movistar).'),
    ('Periférico',
     'Dispositivo que se conecta a una computadora para ampliar sus funciones: monitor, mouse, teclado, cargador, etc.'),
    ('Propiedad (Propia)',
     'Modalidad de propiedad en la que el equipo es comprado y pertenece directamente a la empresa, sin contratos de arrendamiento.'),
    ('Razón social',
     'Nombre legal o denominación oficial de la empresa o entidad con la que se firma el contrato de servicio de un celular.'),
    ('Rol',
     'Nivel de acceso de un usuario dentro del sistema. Puede ser Administrador (acceso total) o Usuario (acceso estándar sin gestión de cuentas).'),
    ('S.O. / Sistema Operativo',
     'Software base que administra los recursos de un equipo. Ejemplos: Windows 11, macOS, Android, iOS, iPadOS.'),
    ('Sucursal / Oficina',
     'Ubicación física o unidad de negocio a la que pertenece un empleado. Se usa en los filtros del Dashboard y la lista de empleados.'),
    ('Etiqueta de inventario',
     'Código o número interno asignado por la empresa para identificar un activo dentro de su propio sistema de control. Puede ser diferente al número de serie del fabricante.'),
]

for termino, definicion in glosario:
    # Término en naranja + negrita
    p = doc.add_paragraph()
    run = p.add_run(termino)
    run.font.bold      = True
    run.font.color.rgb = NARANJA
    run.font.size      = Pt(11)
    pPr = p._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '140')
    sp.set(qn('w:after'),  '20')
    pPr.append(sp)

    # Definición con sangría
    p2 = doc.add_paragraph()
    run2 = p2.add_run(definicion)
    run2.font.color.rgb = GRIS_OSC
    run2.font.size      = Pt(11)
    pPr2 = p2._p.get_or_add_pPr()
    sp2  = OxmlElement('w:spacing')
    sp2.set(qn('w:before'), '0')
    sp2.set(qn('w:after'),  '20')
    pPr2.append(sp2)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '440')
    pPr2.append(ind)

    # Línea divisoria sutil
    p3 = doc.add_paragraph()
    r3 = p3.add_run('─' * 90)
    r3.font.color.rgb = RGBColor(0xE8, 0xE8, 0xE8)
    r3.font.size      = Pt(6)
    no_space_para(p3)

# ══════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ══════════════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run('Select Shop  ·  Sistema de Gestión de Activos  ·  Manual de Usuario v1.0  ·  Junio 2026')
    r1.font.color.rgb = GRIS_MED
    r1.font.size      = Pt(8.5)

# ══════════════════════════════════════════════════════════
output = r'C:\Users\Sistemas2\Desktop\Manual_Sistema_Activos_SelectShop.docx'
doc.save(output)
print(f'Manual guardado en: {output}')
